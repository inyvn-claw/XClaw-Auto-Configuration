#!/usr/bin/env python3
"""
论文模板填空脚本 - 为 Writer 智能体提供 Word 文档生成能力

用法:
  python fill_template.py --create-template    # 创建空白模板
  python fill_template.py --fill               # 使用 data.json 填空
  python fill_template.py --fill --data custom.json  # 使用自定义数据文件
"""

import json
import sys
import argparse
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 占位符定义
PLACEHOLDERS = {
    'TITLE': '论文标题',
    'AUTHOR_1': '作者1',
    'AUTHOR_2': '作者2',
    'AUTHOR_3': '作者3',
    'AUTHOR_4': '作者4',
    'ABSTRACT': '摘要内容',
    'KEYWORDS': '关键字',
    'SEC1_INTRO': 'I 引言',
    'SEC2_RELATED': 'II 相关工作',
    'SEC3_OVERVIEW': 'III-A 模型概述',
    'SEC3_MODULE1': 'III-B 第一模块',
    'SEC3_MODULE2': 'III-C 第二模块',
    'SEC3_MODULE3': 'III-D 第三模块',
    'SEC3_REMARKS': 'III-E 方法讨论',
    'SEC4_SETTING': 'IV-A 实验设置',
    'SEC4_PERFORMANCE': 'IV-B 性能对比',
    'SEC4_ABLATION': 'IV-C 消融研究',
    'SEC4_CONVERGENCE': 'IV-D 收敛性分析',
    'SEC4_PARAMETER': 'IV-E 参数敏感性分析',
    'SEC4_CASE': 'IV-F 案例研究',
    'SEC4_SIGNIFICANCE': 'IV-G 显著性测试',
    'SEC5_CONCLUSION': 'V 结论',
    'REFERENCES': '参考文献'
}

def fix_json_quotes(content):
    """
    修复 JSON 内容中的非法引号
    将字符串值内部的未转义双引号替换为中文直角引号
    """
    lines = content.split('\n')
    fixed_count = 0
    
    for i, line in enumerate(lines):
        line_stripped = line.rstrip()
        
        # 只处理以 " 或 ", 结尾的行（JSON 字符串值行）
        if not (line_stripped.endswith('"') or line_stripped.endswith('",')):
            continue
        
        if ':' not in line:
            continue
        
        colon_pos = line.find(':')
        
        # 找到值开始的引号（在冒号之后）
        value_start = line.find('"', colon_pos)
        if value_start == -1:
            continue
        
        # 找到值结束的引号
        end_pos = len(line_stripped)
        if line_stripped.endswith(','):
            end_pos -= 1
        value_end = end_pos - 1
        
        if value_end <= value_start:
            continue
        
        # 提取值内容
        value = line[value_start+1:value_end]
        
        # 检查值内部是否有未转义的引号
        if '"' in value:
            # 将值内部的 " 替换为中文直角引号
            new_value = value.replace('"', '「')
            new_value = new_value.replace('"', '」')
            if new_value != value:
                lines[i] = line[:value_start+1] + new_value + line[value_end:]
                fixed_count += 1
    
    return '\n'.join(lines), fixed_count

def clean_control_chars(content):
    """
    清理 JSON 内容中的控制字符
    移除所有非标准控制字符（保留 TAB、LF、CR）
    """
    # 定义允许的控制字符：TAB(9), LF(10), CR(13)
    allowed_ctrl = {9, 10, 13}
    
    cleaned = []
    removed_count = 0
    for char in content:
        code = ord(char)
        if code < 32 and code not in allowed_ctrl:
            # 替换控制字符为空字符串或空格
            cleaned.append('')
            removed_count += 1
        else:
            cleaned.append(char)
    
    return ''.join(cleaned), removed_count

def clean_json_content(content):
    """
    清理 JSON 内容中的各种问题：
    1. 控制字符
    2. 未转义的反斜杠（LaTeX公式等）
    """
    # 第一步：处理控制字符
    allowed_ctrl = {9, 10, 13}  # TAB, LF, CR
    cleaned = []
    for char in content:
        code = ord(char)
        if code < 32 and code not in allowed_ctrl:
            cleaned.append('')  # 移除控制字符
        else:
            cleaned.append(char)
    content = ''.join(cleaned)
    
    # 第二步：修复未转义的反斜杠（在字符串值中）
    # 找到所有字符串值，将其中的 \ 替换为 \\
    result = []
    in_string = False
    i = 0
    while i < len(content):
        char = content[i]
        
        if char == '"' and (i == 0 or content[i-1] != '\\'):
            in_string = not in_string
            result.append(char)
        elif in_string and char == '\\':
            # 检查下一个字符
            if i + 1 < len(content):
                next_char = content[i + 1]
                # 如果是有效的 JSON 转义字符，保留
                if next_char in '"\\/bfnrtu':
                    result.append(char)
                else:
                    # 无效的转义，将单个 \ 替换为 \\
                    result.append('\\\\')
            else:
                result.append('\\\\')
        else:
            result.append(char)
        i += 1
    
    return ''.join(result)

def load_json_with_fix(filepath):
    """
    加载 JSON 文件，自动修复各种问题
    """
    with open(filepath, 'r', encoding='utf-8') as f:
        content = f.read()
    
    try:
        # 先尝试正常解析
        return json.loads(content)
    except json.JSONDecodeError as e:
        print(f'⚠️  JSON 解析错误: {e}')
        print(f'   位置: 第 {e.lineno} 行, 第 {e.colno} 列')
        print('\n🔄 尝试自动修复...')
        
        # 第一步：清理 JSON 内容（控制字符 + 反斜杠）
        cleaned_content = clean_json_content(content)
        
        # 第二步：修复引号
        fixed_content, quote_count = fix_json_quotes(cleaned_content)
        if quote_count > 0:
            print(f'   修复了 {quote_count} 行中的非法引号')
        
        try:
            data = json.loads(fixed_content)
            print('✅ 自动修复成功！')
            return data
        except json.JSONDecodeError as e2:
            print(f'❌ 自动修复失败: {e2}')
            print('\n💡 手动修复建议:')
            print(f'   1. 打开 {filepath} 文件')
            print(f'   2. 定位到第 {e2.lineno} 行附近')
            print('   3. 检查是否有 LaTeX 公式中的未转义反斜杠')
            print('   4. 将 \\ 替换为 \\\\')
            raise

def create_template(output_path='templates/V1_paper_template.docx'):
    """创建带占位符的 Word 模板"""
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    
    # 标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('{{TITLE}}')
    run.bold = True
    run.font.size = Pt(18)
    title.space_after = Pt(12)
    
    # 作者
    authors = doc.add_paragraph()
    authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    authors.add_run('{{AUTHOR_1}}，{{AUTHOR_2}}，{{AUTHOR_3}}，{{AUTHOR_4}}')
    authors.space_after = Pt(24)
    
    # 摘要
    abstract_heading = doc.add_paragraph()
    run = abstract_heading.add_run('摘要（Abstract）')
    run.bold = True
    abstract_heading.space_after = Pt(6)
    
    abstract = doc.add_paragraph('{{ABSTRACT}}')
    abstract.space_after = Pt(12)
    
    # 关键字
    keywords = doc.add_paragraph()
    keywords.add_run('关键字：').bold = True
    keywords.add_run('{{KEYWORDS}}')
    keywords.space_after = Pt(24)
    
    # I 引言
    sec1_heading = doc.add_paragraph()
    run = sec1_heading.add_run('I 引言（Introduction）')
    run.bold = True
    sec1_heading.space_after = Pt(6)
    
    doc.add_paragraph('{{SEC1_INTRO}}')
    doc.add_paragraph().space_after = Pt(12)
    
    # II 相关工作
    sec2_heading = doc.add_paragraph()
    run = sec2_heading.add_run('II 相关工作（Related Works）')
    run.bold = True
    sec2_heading.space_after = Pt(6)
    
    doc.add_paragraph('{{SEC2_RELATED}}')
    doc.add_paragraph().space_after = Pt(12)
    
    # III 框架
    sec3_heading = doc.add_paragraph()
    run = sec3_heading.add_run('III 框架（Framework）')
    run.bold = True
    sec3_heading.space_after = Pt(6)
    
    # III-A 模型概述
    sec3a_heading = doc.add_paragraph()
    run = sec3a_heading.add_run('A. 模型概述（Framework Overview）')
    run.bold = True
    sec3a_heading.space_after = Pt(6)
    doc.add_paragraph('{{SEC3_OVERVIEW}}')
    doc.add_paragraph().space_after = Pt(6)
    
    # III-B 第一模块
    sec3b_heading = doc.add_paragraph()
    run = sec3b_heading.add_run('B. 第一模块')
    run.bold = True
    sec3b_heading.space_after = Pt(6)
    doc.add_paragraph('{{SEC3_MODULE1}}')
    doc.add_paragraph().space_after = Pt(6)
    
    # III-C 第二模块
    sec3c_heading = doc.add_paragraph()
    run = sec3c_heading.add_run('C. 第二模块')
    run.bold = True
    sec3c_heading.space_after = Pt(6)
    doc.add_paragraph('{{SEC3_MODULE2}}')
    doc.add_paragraph().space_after = Pt(6)
    
    # III-D 第三模块
    sec3d_heading = doc.add_paragraph()
    run = sec3d_heading.add_run('D. 第三模块')
    run.bold = True
    sec3d_heading.space_after = Pt(6)
    doc.add_paragraph('{{SEC3_MODULE3}}')
    doc.add_paragraph().space_after = Pt(6)
    
    # III-E 方法讨论
    sec3e_heading = doc.add_paragraph()
    run = sec3e_heading.add_run('E. 方法讨论（Methodical Remarks）')
    run.bold = True
    sec3e_heading.space_after = Pt(6)
    doc.add_paragraph('{{SEC3_REMARKS}}')
    doc.add_paragraph().space_after = Pt(12)
    
    # IV 实验
    sec4_heading = doc.add_paragraph()
    run = sec4_heading.add_run('IV 实验（Experiments）')
    run.bold = True
    sec4_heading.space_after = Pt(6)
    
    # IV-A 实验设置
    sec4a_heading = doc.add_paragraph()
    run = sec4a_heading.add_run('A. 实验设置（Experimental Setting）')
    run.bold = True
    sec4a_heading.space_after = Pt(6)
    doc.add_paragraph('{{SEC4_SETTING}}')
    doc.add_paragraph().space_after = Pt(6)
    
    # IV-B 性能对比
    sec4b_heading = doc.add_paragraph()
    run = sec4b_heading.add_run('B. 性能对比（Performance Comparison）')
    run.bold = True
    sec4b_heading.space_after = Pt(6)
    doc.add_paragraph('{{SEC4_PERFORMANCE}}')
    doc.add_paragraph().space_after = Pt(6)
    
    # IV-C 消融研究
    sec4c_heading = doc.add_paragraph()
    run = sec4c_heading.add_run('C. 消融研究（Ablation Test）')
    run.bold = True
    sec4c_heading.space_after = Pt(6)
    doc.add_paragraph('{{SEC4_ABLATION}}')
    doc.add_paragraph().space_after = Pt(6)
    
    # IV-D 收敛性分析
    sec4d_heading = doc.add_paragraph()
    run = sec4d_heading.add_run('D. 收敛性分析（Convergence Analysis）')
    run.bold = True
    sec4d_heading.space_after = Pt(6)
    doc.add_paragraph('{{SEC4_CONVERGENCE}}')
    doc.add_paragraph().space_after = Pt(6)
    
    # IV-E 参数敏感性分析
    sec4e_heading = doc.add_paragraph()
    run = sec4e_heading.add_run('E. 参数敏感性分析（Parameter Analysis）')
    run.bold = True
    sec4e_heading.space_after = Pt(6)
    doc.add_paragraph('{{SEC4_PARAMETER}}')
    doc.add_paragraph().space_after = Pt(6)
    
    # IV-F 案例研究
    sec4f_heading = doc.add_paragraph()
    run = sec4f_heading.add_run('F. 案例研究（Case Study）')
    run.bold = True
    sec4f_heading.space_after = Pt(6)
    doc.add_paragraph('{{SEC4_CASE}}')
    doc.add_paragraph().space_after = Pt(6)
    
    # IV-G 显著性测试
    sec4g_heading = doc.add_paragraph()
    run = sec4g_heading.add_run('G. 显著性测试（Significance Test）')
    run.bold = True
    sec4g_heading.space_after = Pt(6)
    doc.add_paragraph('{{SEC4_SIGNIFICANCE}}')
    doc.add_paragraph().space_after = Pt(12)
    
    # V 结论
    sec5_heading = doc.add_paragraph()
    run = sec5_heading.add_run('V 结论（Conclusion）')
    run.bold = True
    sec5_heading.space_after = Pt(6)
    
    doc.add_paragraph('{{SEC5_CONCLUSION}}')
    doc.add_paragraph().space_after = Pt(12)
    
    # 参考文献
    ref_heading = doc.add_paragraph()
    run = ref_heading.add_run('参考文献（Reference）')
    run.bold = True
    ref_heading.space_after = Pt(6)
    
    doc.add_paragraph('{{REFERENCES}}')
    
    # 保存模板
    doc.save(output_path)
    print(f'✅ 模板已创建: {output_path}')
    print(f'📋 包含占位符: {len(PLACEHOLDERS)} 个')
    return output_path

def fill_template(template_path, data, output_path):
    """使用数据填充模板"""
    doc = Document(template_path)
    
    # 标准化数据键名（去掉 {{ 和 }} 如果存在）
    normalized_data = {}
    for key, value in data.items():
        # 如果键名包含 {{ 和 }}，去掉它们
        clean_key = key.strip('{}') if key.startswith('{{') and key.endswith('}}') else key
        normalized_data[clean_key] = value
    
    # 替换所有占位符
    def replace_in_paragraph(para):
        # 合并所有 runs 的完整文本
        full_text = ''.join(run.text for run in para.runs)
        
        # 检查是否包含任何占位符
        has_replacement = False
        for key, value in normalized_data.items():
            placeholder = f'{{{{{key}}}}}'
            if placeholder in full_text:
                full_text = full_text.replace(placeholder, str(value))
                has_replacement = True
        
        # 如果有替换，更新 runs
        if has_replacement and para.runs:
            # 将完整文本放入第一个 run
            para.runs[0].text = full_text
            # 清空其他 runs
            for run in para.runs[1:]:
                run.text = ''
    
    # 处理段落
    for para in doc.paragraphs:
        replace_in_paragraph(para)
    
    # 处理表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_paragraph(para)
    
    # 保存结果
    doc.save(output_path)
    print(f'✅ 论文已生成: {output_path}')
    
    # 统计替换情况
    filled = sum(1 for key in PLACEHOLDERS.keys() if key in normalized_data and normalized_data[key])
    print(f'📊 填充进度: {filled}/{len(PLACEHOLDERS)} ({filled/len(PLACEHOLDERS)*100:.1f}%)')
    
    # 显示未填充的占位符
    empty = [key for key in PLACEHOLDERS.keys() if key not in normalized_data or not normalized_data[key]]
    if empty:
        print(f'⚠️  未填充: {", ".join(empty[:5])}{"..." if len(empty) > 5 else ""}')
    
    return output_path

def create_sample_data():
    """创建示例数据文件"""
    sample_data = {
        'TITLE': '基于多智能体协作的复杂任务求解框架研究',
        'AUTHOR_1': 'XClaw',
        'AUTHOR_2': '作者2',
        'AUTHOR_3': '作者3',
        'AUTHOR_4': '作者4',
        'ABSTRACT': '本文提出了一种新的多智能体协作框架...',
        'KEYWORDS': '多智能体，协作框架，复杂任务求解',
        'SEC1_INTRO': '（引言内容待填写）',
        'SEC2_RELATED': '（相关工作内容待填写）',
        'SEC3_OVERVIEW': '（模型概述待填写）',
        'SEC3_MODULE1': '（第一模块待填写）',
        'SEC3_MODULE2': '（第二模块待填写）',
        'SEC3_MODULE3': '（第三模块待填写）',
        'SEC3_REMARKS': '（方法讨论待填写）',
        'SEC4_SETTING': '（实验设置待填写）',
        'SEC4_PERFORMANCE': '（性能对比待填写）',
        'SEC4_ABLATION': '（消融研究待填写）',
        'SEC4_CONVERGENCE': '（收敛性分析待填写）',
        'SEC4_PARAMETER': '（参数敏感性分析待填写）',
        'SEC4_CASE': '（案例研究待填写）',
        'SEC4_SIGNIFICANCE': '（显著性测试待填写）',
        'SEC5_CONCLUSION': '（结论待填写）',
        'REFERENCES': '（参考文献待填写）'
    }
    
    with open('data.json', 'w', encoding='utf-8') as f:
        json.dump(sample_data, f, ensure_ascii=False, indent=2)
    print('✅ 示例数据已创建: data.json')
    return sample_data

def update_section(doc_path, section_key, new_content):
    """增量更新指定章节"""
    doc = Document(doc_path)
    placeholder = f'{{{{{section_key}}}}}'
    
    found = False
    for para in doc.paragraphs:
        if placeholder in para.text:
            for run in para.runs:
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, new_content)
                    found = True
    
    if found:
        doc.save(doc_path)
        print(f'✅ 已更新章节: {section_key}')
    else:
        print(f'⚠️  未找到占位符: {section_key}')
    
    return found

def main():
    parser = argparse.ArgumentParser(description='论文模板填空工具')
    parser.add_argument('--create-template', action='store_true', help='创建空白模板')
    parser.add_argument('--create-data', action='store_true', help='创建示例数据文件')
    parser.add_argument('--fill', action='store_true', help='填充模板')
    parser.add_argument('--data', default='data.json', help='数据文件路径 (默认: data.json)')
    parser.add_argument('--template', default='templates/V1_paper_template.docx', help='模板路径')
    parser.add_argument('--output', default='output/paper.docx', help='输出路径')
    parser.add_argument('--update', help='更新指定章节 (如: SEC1_INTRO)')
    parser.add_argument('--content', help='更新的内容')
    
    args = parser.parse_args()
    
    if args.create_template:
        create_template(args.template)
    
    elif args.create_data:
        create_sample_data()
    
    elif args.fill:
        # 读取数据（带容错处理）
        try:
            data = load_json_with_fix(args.data)
        except FileNotFoundError:
            print(f'❌ 数据文件不存在: {args.data}')
            print('💡 请先运行: python fill_template.py --create-data')
            return 1
        except json.JSONDecodeError:
            return 1
        
        fill_template(args.template, data, args.output)
    
    elif args.update and args.content:
        update_section(args.output, args.update, args.content)
    
    else:
        parser.print_help()
    
    return 0

if __name__ == '__main__':
    sys.exit(main())
